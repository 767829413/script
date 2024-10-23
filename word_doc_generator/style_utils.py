from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx import Document


# 设置表格字体大小和字体名称的函数
def set_table_style(table,
                    font_size,
                    font_name,
                    alignment=WD_ALIGN_VERTICAL.CENTER):
    for row in table.rows:
        row.height = Pt(20)  # 设置行高，可以根据需要调整
        for cell in row.cells:
            cell.vertical_alignment = alignment
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = font_name
                    # 设置中文字体
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # 应用表格样式
    table.style = 'Table Grid'
    # 设置第一行作为表头
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True


def add_empty_paragraphs(doc: Document, n: int) -> None:
    """
    向文档中添加n个空白段落。

    :param doc: Document对象，表示要操作的Word文档
    :param n: int，要添加的空白段落数量
    :return: None
    """
    for _ in range(n):
        doc.add_paragraph()


def add_heading_with_style(doc,
                           text,
                           level,
                           font_name='宋体',
                           font_size=14,
                           space_after=12):
    """
    添加带有特定样式的标题。

    :param doc: Document对象
    :param text: 标题文本
    :param level: 标题级别（1-9）
    :param font_name: 字体名称
    :param font_size: 字体大小（磅值）
    :param space_after: 段落后的间距（磅值）
    :return: 返回添加的段落对象
    """
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    heading.paragraph_format.space_after = Pt(space_after)
    return heading
