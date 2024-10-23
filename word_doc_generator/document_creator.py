from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from .style_utils import add_heading_with_style, set_table_style, add_empty_paragraphs


def create_table_word_document(data):
    # 创建 Word 文档
    doc = Document()

    # 添加正文，宋体，居中，字号一
    heading = add_heading_with_style(doc,
                                     "融课服务端数据表",
                                     0,
                                     font_size=48,
                                     space_after=24)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 添加标题2，名称是表清单
    add_heading_with_style(
        doc,
        "表清单",
        2,
    )

    # 添加表格，行数根据数据动态设置
    table = doc.add_table(rows=len(data['tables']) + 1,
                          cols=4)  # +1 for the header row
    table.style = 'Table Grid'  # 应用网格样式
    table.autofit = True  # 启用自动调整大小

    # 设置表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '数据表'
    hdr_cells[2].text = '名称'
    hdr_cells[3].text = '备注'

    # 插入数据
    for i, row_data in enumerate(data['tables'], start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = str(i)  # 将整数转换为字符串
        row_cells[1].text = row_data["tableName"]
        row_cells[2].text = row_data["tableComment"]
        row_cells[3].text = ""

    # 设置表索引表格字体大小和字体名称
    set_table_style(table, 10, '宋体')  # 设置字体大小为10磅，字体为宋体

    # 添加标题2，名称是表字段详细
    add_heading_with_style(doc, "表字段详细", 2)

    # 动态添加标题3
    for tableInfo in data['tableInfos']:
        title = tableInfo['tableName'] + " [" + tableInfo['tableComment'] + "]"
        add_heading_with_style(doc, title, 3, font_size=14, space_after=8)

        # 添加表字段表格
        table2 = doc.add_table(rows=len(tableInfo['tableColumns']) + 1, cols=7)
        table2.style = 'Table Grid'  # 应用网格样式
        table2.autofit = True  # 启用自动调整大小
        hdr_cells2 = table2.rows[0].cells
        hdr_cells2[0].text = '字段'
        hdr_cells2[1].text = '数据类型'
        hdr_cells2[2].text = '长度'
        hdr_cells2[3].text = '是否主键'
        hdr_cells2[4].text = '是否为空'
        hdr_cells2[5].text = '默认值'
        hdr_cells2[6].text = '备注'

        # 插入数据
        for i, tableColumn in enumerate(tableInfo['tableColumns'], start=1):
            row_cells = table2.rows[i].cells
            row_cells[0].text = tableColumn["columnName"]
            row_cells[1].text = tableColumn["columnType"]
            row_cells[2].text = str(tableColumn["columnLength"])
            row_cells[3].text = tableColumn["isPrimary"]
            row_cells[4].text = tableColumn["isNullable"]
            row_cells[5].text = tableColumn["columnDefault"] or ""
            row_cells[6].text = tableColumn["columnComment"]

        # 设置表索引表格字体大小和字体名称
        set_table_style(table2, 10, '宋体')  # 设置字体大小为10磅，字体为宋体

        add_empty_paragraphs(doc, 3)

        # 添加表索引表格
        table3 = doc.add_table(rows=len(tableInfo['tableIndexs']) + 1, cols=5)
        table3.style = 'Table Grid'  # 应用网格样式
        table3.autofit = True  # 启用自动调整大小
        hdr_cells3 = table3.rows[0].cells
        hdr_cells3[0].text = '索引名称'
        hdr_cells3[1].text = '索引字段'
        hdr_cells3[2].text = '是否唯一'
        hdr_cells3[3].text = '索引类型'
        hdr_cells3[4].text = '备注'

        # 插入数据
        for i, tableIndex in enumerate(tableInfo['tableIndexs'], start=1):
            row_cells = table3.rows[i].cells
            row_cells[0].text = tableIndex["indexName"]
            row_cells[1].text = tableIndex["columnName"]
            row_cells[2].text = tableIndex["isUnique"]
            row_cells[3].text = tableIndex["indexType"]
            row_cells[4].text = tableIndex["indexComment"]
        # 设置表索引表格字体大小和字体名称
        set_table_style(table3, 10, '宋体')  # 设置字体大小为10磅，字体为宋体

        add_empty_paragraphs(doc, 3)

    return doc
